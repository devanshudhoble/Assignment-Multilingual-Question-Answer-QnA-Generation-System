"""
Document Parser Module
======================
Handles text extraction from PDF, DOCX, and TXT files.
Supports documents written in English, Hindi, Marathi, or mixed languages.
"""

import io
from PyPDF2 import PdfReader
from docx import Document


def parse_pdf(uploaded_file) -> str:
    """
    Extract text from a PDF file.
    
    Args:
        uploaded_file: A file-like object (Streamlit UploadedFile or BytesIO).
    
    Returns:
        Extracted text as a single string.
    """
    try:
        reader = PdfReader(uploaded_file)
        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text.strip())
        return "\n\n".join(text_parts)
    except Exception as e:
        raise ValueError(f"Error reading PDF file: {str(e)}")


def parse_docx(uploaded_file) -> str:
    """
    Extract text from a DOCX (Microsoft Word) file.
    
    Args:
        uploaded_file: A file-like object (Streamlit UploadedFile or BytesIO).
    
    Returns:
        Extracted text as a single string.
    """
    try:
        doc = Document(uploaded_file)
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        return "\n\n".join(text_parts)
    except Exception as e:
        raise ValueError(f"Error reading DOCX file: {str(e)}")


def parse_txt(uploaded_file) -> str:
    """
    Extract text from a plain TXT file.
    
    Args:
        uploaded_file: A file-like object (Streamlit UploadedFile or BytesIO).
    
    Returns:
        Extracted text as a single string.
    """
    try:
        content = uploaded_file.read()
        if isinstance(content, bytes):
            # Try UTF-8 first, then fall back to other encodings
            for encoding in ["utf-8", "utf-16", "latin-1", "cp1252"]:
                try:
                    return content.decode(encoding)
                except (UnicodeDecodeError, LookupError):
                    continue
            return content.decode("utf-8", errors="replace")
        return content
    except Exception as e:
        raise ValueError(f"Error reading TXT file: {str(e)}")


def extract_text(uploaded_file, filename: str) -> str:
    """
    Dispatcher function that routes to the correct parser based on file extension.
    
    Args:
        uploaded_file: A file-like object (Streamlit UploadedFile or BytesIO).
        filename: The name of the uploaded file (used to detect extension).
    
    Returns:
        Extracted text as a single string.
    
    Raises:
        ValueError: If the file format is not supported.
    """
    filename_lower = filename.lower()
    
    if filename_lower.endswith(".pdf"):
        return parse_pdf(uploaded_file)
    elif filename_lower.endswith(".docx"):
        return parse_docx(uploaded_file)
    elif filename_lower.endswith(".txt"):
        return parse_txt(uploaded_file)
    else:
        raise ValueError(
            f"Unsupported file format: {filename}. "
            "Please upload a .pdf, .docx, or .txt file."
        )
