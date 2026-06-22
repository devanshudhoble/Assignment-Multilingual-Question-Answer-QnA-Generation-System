"""
Generate sample test files in PDF, DOCX, and TXT formats
for testing the Multilingual QnA Generation System.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF
import os

OUTPUT_DIR = "sample_files"
os.makedirs(OUTPUT_DIR, exist_ok=True)

# ─────────────────────────────────────────────────────────
# CONTENT FOR EACH FILE
# ─────────────────────────────────────────────────────────

TXT_CONTENT = """Indian Space Research Organisation (ISRO)

The Indian Space Research Organisation (ISRO) is the national space agency of India, headquartered in Bengaluru, Karnataka. It was founded on August 15, 1969, by Dr. Vikram Sarabhai, who is widely regarded as the father of the Indian space programme. ISRO operates under the Department of Space (DOS), which is directly overseen by the Prime Minister of India.

History and Evolution
India's space journey began in 1962 when the Indian National Committee for Space Research (INCOSPAR) was established by Jawaharlal Nehru. The first sounding rocket was launched from Thumba Equatorial Rocket Launching Station (TERLS) in Kerala on November 21, 1963. ISRO was officially formed in 1969, replacing INCOSPAR.

The first Indian satellite, Aryabhata, was launched on April 19, 1975, by the Soviet Union. This was followed by the Bhaskara-I satellite in 1979. India's first indigenous launch vehicle, SLV-3, successfully placed the Rohini satellite in orbit on July 18, 1980, making India the sixth nation to achieve this capability.

Major Missions and Achievements
Chandrayaan-1, launched on October 22, 2008, was India's first mission to the Moon. It made the groundbreaking discovery of water molecules on the lunar surface. The mission carried 11 scientific instruments from India, USA, UK, Germany, Sweden, and Bulgaria.

The Mars Orbiter Mission (Mangalyaan), launched on November 5, 2013, made India the first Asian nation to reach Martian orbit and the first nation in the world to do so on its maiden attempt. The mission cost approximately $74 million, making it the least expensive Mars mission at the time.

Chandrayaan-3, launched on July 14, 2023, successfully achieved a soft landing near the lunar south pole on August 23, 2023. This made India the fourth country to successfully land on the Moon and the first to land near the south polar region.

Launch Vehicles
ISRO has developed several launch vehicles including the Polar Satellite Launch Vehicle (PSLV) and the Geosynchronous Satellite Launch Vehicle (GSLV). The PSLV is known for its reliability and versatility, having completed over 50 successful missions. In February 2017, PSLV-C37 set a world record by deploying 104 satellites in a single mission.

The GSLV Mark III, also known as LVM3, is India's heaviest launch vehicle capable of placing 4-tonne class satellites into Geosynchronous Transfer Orbit (GTO). It was used for the Chandrayaan-3 mission.

Satellite Systems
ISRO operates the Indian National Satellite System (INSAT) for communication, television broadcasting, and meteorology services. The Indian Remote Sensing (IRS) satellite system is one of the largest civilian remote sensing satellite constellations in the world, providing data for agriculture, water resources, urban planning, and disaster management.

The Navigation with Indian Constellation (NavIC), also known as the Indian Regional Navigation Satellite System (IRNSS), provides accurate position information to users in India and the surrounding region up to 1,500 km from the Indian boundary.

Future Plans
ISRO is currently working on several ambitious projects including Gaganyaan, India's first crewed spaceflight programme, expected to send Indian astronauts to space. The agency is also planning Chandrayaan-4 for lunar sample return, Shukrayaan-1 for Venus exploration, and an Indian Space Station by 2035.

Budget and Global Standing
ISRO's annual budget for 2023-24 was approximately $1.5 billion. Despite having one of the smaller budgets among major space agencies, ISRO is recognized worldwide for its cost-effective missions. The agency has launched satellites for 34 countries, generating significant revenue and establishing India as a reliable launch service provider.
"""

DOCX_CONTENT = {
    "title": "Climate Change and Its Global Impact",
    "sections": [
        {
            "heading": "Introduction to Climate Change",
            "text": "Climate change refers to long-term shifts in temperatures and weather patterns across the globe. While some of these shifts may be natural, since the 1800s, human activities have been the primary driver of climate change, mainly due to the burning of fossil fuels such as coal, oil, and natural gas. These activities produce greenhouse gases that act like a blanket wrapped around the Earth, trapping the Sun's heat and raising global temperatures. The Earth's average surface temperature has risen by approximately 1.1 degrees Celsius since the pre-industrial era."
        },
        {
            "heading": "Causes of Climate Change",
            "text": "The primary cause of climate change is the emission of greenhouse gases (GHGs) into the atmosphere. Carbon dioxide (CO2) is the most significant GHG, accounting for about 76% of total emissions. It is released primarily through the burning of fossil fuels for electricity, heat, and transportation. Methane (CH4), which accounts for about 16% of emissions, comes from agriculture, waste management, and energy production. Nitrous oxide (N2O) makes up about 6% and is released from agricultural practices and industrial processes. Deforestation contributes to climate change by reducing the number of trees that absorb CO2 from the atmosphere. The industrial revolution, which began in the mid-18th century, marked a significant increase in greenhouse gas emissions."
        },
        {
            "heading": "Effects on the Environment",
            "text": "Rising global temperatures are causing widespread environmental changes. Arctic sea ice is declining at a rate of 13% per decade. Glaciers worldwide are retreating, contributing to rising sea levels, which have increased by about 20 centimeters since 1900. Extreme weather events, including hurricanes, droughts, floods, and heatwaves, are becoming more frequent and intense. Ocean acidification, caused by the absorption of excess CO2, threatens marine ecosystems and coral reefs. The Great Barrier Reef in Australia has experienced multiple mass bleaching events. Biodiversity loss is accelerating as species struggle to adapt to rapidly changing conditions."
        },
        {
            "heading": "Impact on Human Society",
            "text": "Climate change affects human health through increased air pollution, the spread of vector-borne diseases, and heat-related illnesses. Agriculture is heavily impacted, with changing rainfall patterns and extreme weather threatening food security for millions. The World Bank estimates that climate change could push 132 million people into extreme poverty by 2030. Small island nations and coastal communities face existential threats from rising sea levels. Climate migration is increasing, with the Internal Displacement Monitoring Centre reporting 30 million climate-related displacements annually."
        },
        {
            "heading": "The Paris Agreement",
            "text": "The Paris Agreement, adopted in December 2015 at COP21, is a landmark international treaty on climate change. It was signed by 196 parties and aims to limit global warming to well below 2 degrees Celsius above pre-industrial levels, with efforts to limit it to 1.5 degrees. Each country submits Nationally Determined Contributions (NDCs) outlining their climate action plans. India committed to reducing the emissions intensity of its GDP by 45% by 2030 compared to 2005 levels, and achieving net-zero emissions by 2070."
        },
        {
            "heading": "Solutions and Mitigation",
            "text": "Transitioning to renewable energy sources such as solar, wind, and hydroelectric power is essential for reducing greenhouse gas emissions. Energy efficiency improvements in buildings, transportation, and industry can significantly reduce energy consumption. Reforestation and afforestation help absorb CO2 from the atmosphere. Electric vehicles are gaining popularity as a cleaner alternative to internal combustion engines. Carbon capture and storage (CCS) technology aims to capture CO2 emissions from power plants and industrial processes. Individual actions such as reducing energy consumption, minimizing waste, choosing sustainable transportation, and adopting plant-based diets also contribute to climate action."
        }
    ]
}

PDF_CONTENT = {
    "title": "The History and Evolution of the Internet",
    "sections": [
        ("Origins of the Internet",
         "The Internet as we know it today has its roots in a project called ARPANET (Advanced Research Projects Agency Network), which was developed by the United States Department of Defense in the late 1960s. ARPANET was designed to enable communication between research institutions and to create a network that could survive partial outages. The first message was sent over ARPANET on October 29, 1969, from a computer at UCLA to one at Stanford Research Institute. The message was supposed to be 'LOGIN,' but the system crashed after transmitting just the letters 'L' and 'O.'"),

        ("Development of TCP/IP",
         "In the 1970s, Vinton Cerf and Bob Kahn developed the Transmission Control Protocol (TCP) and Internet Protocol (IP), which became the foundational communication protocols of the Internet. TCP/IP was adopted as the standard protocol for ARPANET on January 1, 1983, which is often considered the official birth date of the Internet. These protocols enabled different computer networks to interconnect and communicate with each other, creating a 'network of networks.'"),

        ("The World Wide Web",
         "In 1989, Sir Tim Berners-Lee, a British computer scientist working at CERN (European Organization for Nuclear Research) in Switzerland, proposed a system for sharing information using hypertext. He developed three fundamental technologies: HTML (HyperText Markup Language), URI (Uniform Resource Identifier), and HTTP (HyperText Transfer Protocol). The first website went live on August 6, 1991, and was dedicated to information about the World Wide Web project itself. Berners-Lee made the World Wide Web freely available to everyone, with no patent and no royalties."),

        ("The Rise of Web Browsers and Search Engines",
         "The first widely popular web browser, Mosaic, was released in 1993 by the National Center for Supercomputing Applications (NCSA). It was followed by Netscape Navigator in 1994, which dominated the early browser market. Microsoft entered the browser market with Internet Explorer in 1995, leading to the famous 'browser wars.' Early search engines like AltaVista (1995), Yahoo (1994), and Ask Jeeves (1996) helped users navigate the growing World Wide Web. Google, founded by Larry Page and Sergey Brin in September 1998, revolutionized search with its PageRank algorithm, which ranked websites based on the number and quality of links pointing to them."),

        ("The Social Media Revolution",
         "The mid-2000s saw the rise of social media platforms that transformed how people communicate and share information online. Facebook was launched by Mark Zuckerberg in February 2004 from his Harvard University dorm room and has grown to over 3 billion monthly active users. YouTube, founded in 2005, became the world's largest video-sharing platform. Twitter (now X), launched in 2006, popularized microblogging. Instagram (2010) and TikTok (2016) further changed the landscape of online communication with their focus on visual and short-form video content."),

        ("The Mobile Internet Era",
         "The launch of Apple's iPhone in 2007 and the subsequent Android platform by Google marked the beginning of the mobile Internet era. By 2023, mobile devices accounted for approximately 60% of all global web traffic. Mobile applications (apps) have created entirely new industries and transformed existing ones, from ride-sharing (Uber, Ola) to food delivery (Swiggy, Zomato) to digital payments (Google Pay, PhonePe). The rollout of 4G LTE networks and the ongoing deployment of 5G technology have enabled faster data speeds and lower latency, supporting applications like video streaming, cloud gaming, and the Internet of Things (IoT)."),

        ("Internet Today and Future",
         "As of 2024, there are approximately 5.4 billion Internet users worldwide, representing about 67% of the global population. The Internet has become essential infrastructure for education, commerce, entertainment, and governance. Emerging technologies like artificial intelligence, blockchain, quantum computing, and Web3 are shaping the next evolution of the Internet. The concept of the metaverse, a persistent virtual world, represents one possible future of the Internet. Challenges including digital divide, cybersecurity threats, data privacy, and misinformation continue to be major concerns for the global Internet community.")
    ]
}


def create_txt_file():
    """Create a sample .txt file about ISRO."""
    filepath = os.path.join(OUTPUT_DIR, "ISRO_Space_Programme.txt")
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(TXT_CONTENT.strip())
    print(f"  Created: {filepath}")
    return filepath


def create_docx_file():
    """Create a sample .docx file about Climate Change."""
    filepath = os.path.join(OUTPUT_DIR, "Climate_Change_Report.docx")
    doc = Document()

    # Title
    title = doc.add_heading(DOCX_CONTENT["title"], level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("A Comprehensive Overview of Causes, Effects, and Solutions")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(100, 100, 100)
    run.font.italic = True

    doc.add_paragraph()  # spacer

    for section in DOCX_CONTENT["sections"]:
        doc.add_heading(section["heading"], level=1)
        para = doc.add_paragraph(section["text"])
        para.paragraph_format.space_after = Pt(8)
        for run in para.runs:
            run.font.size = Pt(11)

    doc.save(filepath)
    print(f"  Created: {filepath}")
    return filepath


def create_pdf_file():
    """Create a sample .pdf file about the Internet."""
    filepath = os.path.join(OUTPUT_DIR, "History_of_Internet.pdf")

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()

    # Title
    pdf.set_font("Helvetica", "B", 22)
    pdf.cell(0, 15, PDF_CONTENT["title"], new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(5)

    # Subtitle
    pdf.set_font("Helvetica", "I", 11)
    pdf.set_text_color(120, 120, 120)
    pdf.cell(0, 8, "From ARPANET to the Modern Digital World", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.set_text_color(0, 0, 0)
    pdf.ln(10)

    # Sections
    for heading, text in PDF_CONTENT["sections"]:
        pdf.set_font("Helvetica", "B", 14)
        pdf.set_text_color(30, 60, 120)
        pdf.cell(0, 10, heading, new_x="LMARGIN", new_y="NEXT")
        pdf.set_text_color(0, 0, 0)
        pdf.ln(2)

        pdf.set_font("Helvetica", "", 11)
        pdf.multi_cell(0, 6, text)
        pdf.ln(6)

    pdf.output(filepath)
    print(f"  Created: {filepath}")
    return filepath


if __name__ == "__main__":
    print("\nGenerating sample test files...\n")
    create_txt_file()
    create_docx_file()
    create_pdf_file()
    print(f"\nAll files saved to: {OUTPUT_DIR}/")
    print("You can now upload these in the Streamlit app to test!")
